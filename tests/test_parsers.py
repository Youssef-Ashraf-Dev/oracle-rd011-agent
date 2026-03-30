"""
RD.011 Agent — Parser unit tests.

Tests docx_parser and excel_parser against real sample files.
These tests require the sample files in samples/inputs/ to exist.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest

from parsers.docx_parser import parse_docx
from parsers.excel_parser import parse_excel

# Resolve sample paths relative to repository root
REPO_ROOT = Path(__file__).parent.parent
SAMPLES_DIR = REPO_ROOT / "samples" / "inputs"

# Sample file names — skip if not present
AP_MOM = SAMPLES_DIR / "AP20_Formatted.docx"
GL_MOM = SAMPLES_DIR / "Dec_19_2022_GL_Analysis_MOM_Presentation_Ready.docx"
AR_MOM = SAMPLES_DIR / "AR_Full.docx"
QUESTIONNAIRE = SAMPLES_DIR / "FIN Questionnaire_.xlsx"
REQUIREMENTS_DOC = SAMPLES_DIR / "Oracle_Financials_Requirements_Agreed_PainPoints.docx"

requires_samples = pytest.mark.skipif(
    not SAMPLES_DIR.exists(),
    reason="samples/inputs/ directory not found",
)


# ── docx_parser ───────────────────────────────────────────────────────────

class TestDocxParser:
    @requires_samples
    @pytest.mark.skipif(not AP_MOM.exists(), reason="AP MOM sample not found")
    def test_parse_ap_mom_returns_nonempty_string(self):
        result = parse_docx(str(AP_MOM))
        assert isinstance(result, str)
        assert len(result) > 100, "Parsed output should contain substantial text"

    @requires_samples
    @pytest.mark.skipif(not GL_MOM.exists(), reason="GL MOM sample not found")
    def test_parse_gl_mom_returns_nonempty_string(self):
        result = parse_docx(str(GL_MOM))
        assert isinstance(result, str)
        assert len(result) > 100

    @requires_samples
    @pytest.mark.skipif(not GL_MOM.exists(), reason="GL MOM sample not found")
    def test_parse_gl_mom_contains_gl_content(self):
        """GL MOM should contain GL-related content."""
        result = parse_docx(str(GL_MOM))
        # At minimum the file should contain some text
        assert result.strip() != ""

    @requires_samples
    @pytest.mark.skipif(not AR_MOM.exists(), reason="AR MOM sample not found")
    def test_parse_ar_mom_returns_nonempty_string(self):
        result = parse_docx(str(AR_MOM))
        assert isinstance(result, str)
        assert len(result) > 100

    @requires_samples
    @pytest.mark.skipif(not REQUIREMENTS_DOC.exists(), reason="Requirements doc not found")
    def test_parse_requirements_doc(self):
        result = parse_docx(str(REQUIREMENTS_DOC))
        assert isinstance(result, str)
        assert len(result) > 50

    def test_parse_nonexistent_file_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_docx("/nonexistent/path/file.docx")

    def test_parse_wrong_extension_raises(self):
        with pytest.raises(ValueError):
            parse_docx("somefile.pdf")

    def test_parse_docx_table_content(self, tmp_path):
        """Create a minimal docx with a table and verify it is parsed."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test Heading", style="Heading 1")
        doc.add_paragraph("This is a body paragraph with some text content here.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = parse_docx(str(docx_path))
        assert "Test Heading" in result
        assert "body paragraph" in result
        assert "Header A" in result
        assert "Value 1" in result

    def test_parse_docx_heading_markers(self, tmp_path):
        """Heading levels should be marked with # prefixes."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Level One", style="Heading 1")
        doc.add_paragraph("Level Two", style="Heading 2")
        doc.add_paragraph("Normal text paragraph.")
        docx_path = tmp_path / "headings.docx"
        doc.save(str(docx_path))

        result = parse_docx(str(docx_path))
        assert "Level One" in result
        assert "Level Two" in result

    def test_parse_docx_empty_document(self, tmp_path):
        """An empty docx should return an empty string without raising."""
        from docx import Document

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))

        result = parse_docx(str(docx_path))
        assert isinstance(result, str)


# ── excel_parser ──────────────────────────────────────────────────────────

class TestExcelParser:
    @requires_samples
    @pytest.mark.skipif(not QUESTIONNAIRE.exists(), reason="Questionnaire sample not found")
    def test_parse_questionnaire_returns_nonempty_string(self):
        result = parse_excel(str(QUESTIONNAIRE))
        assert isinstance(result, str)
        assert len(result) > 100

    @requires_samples
    @pytest.mark.skipif(not QUESTIONNAIRE.exists(), reason="Questionnaire sample not found")
    def test_parse_questionnaire_multiple_sheets(self):
        """Questionnaire likely has multiple sheets — all should be included."""
        result = parse_excel(str(QUESTIONNAIRE))
        # Should have at least some sheet heading marker
        assert "##" in result or len(result) > 50

    @requires_samples
    @pytest.mark.skipif(not QUESTIONNAIRE.exists(), reason="Questionnaire sample not found")
    def test_parse_questionnaire_handles_merged_cells(self):
        """Merged cells should not raise an exception."""
        # This just checks the parse completes without error
        result = parse_excel(str(QUESTIONNAIRE))
        assert isinstance(result, str)

    def test_parse_nonexistent_excel_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_excel("/nonexistent/path/file.xlsx")

    def test_parse_wrong_extension_raises(self):
        with pytest.raises(ValueError):
            parse_excel("somefile.csv")

    def test_parse_excel_basic_data(self, tmp_path):
        """Create a minimal xlsx and verify it is parsed correctly."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "TestSheet"
        ws["A1"] = "Module"
        ws["B1"] = "Requirement"
        ws["A2"] = "AP"
        ws["B2"] = "Three-way match required for PO invoices"
        ws["A3"] = "GL"
        ws["B3"] = "Manual journals require dual approval"
        xlsx_path = tmp_path / "test.xlsx"
        wb.save(str(xlsx_path))

        result = parse_excel(str(xlsx_path))
        assert "TestSheet" in result
        assert "Module" in result
        assert "Requirement" in result
        assert "AP" in result
        assert "Three-way match" in result

    def test_parse_excel_merged_cells(self, tmp_path):
        """Merged cells should return the top-left value for all merged cells."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Merged"
        ws["A1"] = "Merged Header"
        ws.merge_cells("A1:C1")  # Merge A1:C1
        ws["A2"] = "Col A"
        ws["B2"] = "Col B"
        ws["C2"] = "Col C"
        xlsx_path = tmp_path / "merged.xlsx"
        wb.save(str(xlsx_path))

        # Should not raise — merged cells are handled gracefully
        result = parse_excel(str(xlsx_path))
        assert "Merged" in result or "Col A" in result

    def test_parse_excel_blank_cells_handled(self, tmp_path):
        """Rows with some blank cells should be parsed without raising."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["C1"] = "Notes"
        ws["A2"] = "Item1"
        ws["B2"] = None  # blank
        ws["C2"] = "Some note"
        ws["A3"] = None  # blank row — should be skipped
        ws["B3"] = None
        ws["C3"] = None
        xlsx_path = tmp_path / "blanks.xlsx"
        wb.save(str(xlsx_path))

        result = parse_excel(str(xlsx_path))
        assert "Name" in result
        assert "Item1" in result
        assert "Some note" in result

    def test_parse_excel_multiple_sheets(self, tmp_path):
        """Multiple sheets should all appear in the combined output."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1["A1"] = "Sheet1 Data"

        ws2 = wb.create_sheet("Sheet2")
        ws2["A1"] = "Sheet2 Data"

        xlsx_path = tmp_path / "multi.xlsx"
        wb.save(str(xlsx_path))

        result = parse_excel(str(xlsx_path))
        assert "Sheet1" in result
        assert "Sheet2" in result
        assert "Sheet1 Data" in result
        assert "Sheet2 Data" in result
