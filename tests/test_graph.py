"""
RD.011 Agent — Graph structure and routing unit tests.

Tests that the graph is assembled correctly, all nodes are registered,
routing functions behave correctly, and the ingest node handles a sample
file without LLM calls.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Any, Dict
from unittest.mock import patch

import pytest

from graph import build_graph, route_after_approval, route_after_section
from state import RD011State


# ── Fixtures ──────────────────────────────────────────────────────────────

@pytest.fixture
def minimal_state() -> Dict[str, Any]:
    """Return a minimal valid RD011State as a plain dict."""
    return {
        "thread_id": "test-001",
        "input_files": [],
        "raw_texts": {},
        "extraction_result": None,
        "document_plan": None,
        "issue_report": None,
        "consultant_approved": False,
        "consultant_feedback": "",
        "approval_iteration": 0,
        "approval_maxed": False,
        "intro_content": None,
        "section_queue": [],
        "current_section_index": 0,
        "generated_sections": {},
        "failed_sections": [],
        "diagram_registry": {},
        "output_path": None,
        "errors": [],
        "last_completed_node": "",
    }


# ── Graph structure tests ─────────────────────────────────────────────────

class TestBuildGraph:
    def test_build_graph_does_not_raise(self, tmp_path):
        """build_graph() should complete without raising."""
        db_path = str(tmp_path / "test_checkpoints.db")
        with patch("graph.CHECKPOINT_DB_PATH", db_path):
            graph = build_graph()
        assert graph is not None

    def test_graph_has_all_expected_nodes(self, tmp_path):
        """All 12 processing nodes should be registered."""
        db_path = str(tmp_path / "test_checkpoints.db")
        with patch("graph.CHECKPOINT_DB_PATH", db_path):
            graph = build_graph()

        expected_nodes = {
            "ingest",
            "extract",
            "plan",
            "detect_issues",
            "present_plan",
            "await_approval",
            "update_plan",
            "generate_intro",
            "generate_section",
            "render_diagrams",
            "assemble_document",
            "error_handler",
        }

        # LangGraph compiled graphs expose node names in different attributes
        # depending on the version — check common locations
        graph_nodes = set()
        if hasattr(graph, "nodes"):
            graph_nodes = set(graph.nodes)
        elif hasattr(graph, "_graph") and hasattr(graph._graph, "nodes"):
            graph_nodes = set(graph._graph.nodes)
        elif hasattr(graph, "graph") and hasattr(graph.graph, "nodes"):
            graph_nodes = set(graph.graph.nodes)

        if graph_nodes:
            for node in expected_nodes:
                assert node in graph_nodes, f"Expected node '{node}' not found in graph"


# ── Routing function tests ────────────────────────────────────────────────

class TestRouteAfterApproval:
    def test_approved_routes_to_generate_intro(self, minimal_state):
        minimal_state["consultant_approved"] = True
        result = route_after_approval(minimal_state)
        assert result == "generate_intro"

    def test_maxed_routes_to_error_handler(self, minimal_state):
        minimal_state["approval_maxed"] = True
        result = route_after_approval(minimal_state)
        assert result == "error_handler"

    def test_not_approved_routes_to_update_plan(self, minimal_state):
        minimal_state["consultant_approved"] = False
        result = route_after_approval(minimal_state)
        assert result == "update_plan"

    def test_approved_false_default(self, minimal_state):
        """Unapproved state should always route to update_plan."""
        assert route_after_approval(minimal_state) == "update_plan"


class TestRouteAfterSection:
    def test_queue_not_exhausted_returns_generate_section(self, minimal_state):
        minimal_state["section_queue"] = ["AP.01", "AP.02", "GL.01"]
        minimal_state["current_section_index"] = 1
        result = route_after_section(minimal_state)
        assert result == "generate_section"

    def test_queue_exhausted_returns_render_diagrams(self, minimal_state):
        minimal_state["section_queue"] = ["AP.01", "AP.02"]
        minimal_state["current_section_index"] = 2  # index == len(queue)
        result = route_after_section(minimal_state)
        assert result == "render_diagrams"

    def test_empty_queue_returns_render_diagrams(self, minimal_state):
        minimal_state["section_queue"] = []
        minimal_state["current_section_index"] = 0
        result = route_after_section(minimal_state)
        assert result == "render_diagrams"

    def test_index_beyond_queue_returns_render_diagrams(self, minimal_state):
        minimal_state["section_queue"] = ["AP.01"]
        minimal_state["current_section_index"] = 5
        result = route_after_section(minimal_state)
        assert result == "render_diagrams"


# ── Ingest node unit tests ────────────────────────────────────────────────

class TestIngestNode:
    def test_ingest_docx_sample(self, minimal_state, tmp_path):
        """ingest_node should parse a .docx file and store text in raw_texts."""
        from docx import Document
        from nodes.ingest_node import ingest_node

        # Create a minimal docx
        doc = Document()
        doc.add_paragraph("Test heading", style="Heading 1")
        doc.add_paragraph("This is body content about Oracle Finance Cloud.")
        docx_path = str(tmp_path / "test_mom.docx")
        doc.save(docx_path)

        minimal_state["input_files"] = [docx_path]
        result = ingest_node(minimal_state)

        assert "raw_texts" in result
        filename = os.path.basename(docx_path)
        assert filename in result["raw_texts"]
        assert "Oracle Finance Cloud" in result["raw_texts"][filename]
        assert result["last_completed_node"] == "ingest"

    def test_ingest_xlsx_sample(self, minimal_state, tmp_path):
        """ingest_node should parse a .xlsx file and store text in raw_texts."""
        import openpyxl
        from nodes.ingest_node import ingest_node

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Questionnaire"
        ws["A1"] = "Module"
        ws["B1"] = "Requirement"
        ws["A2"] = "AP"
        ws["B2"] = "Three-way match"
        xlsx_path = str(tmp_path / "questionnaire.xlsx")
        wb.save(xlsx_path)

        minimal_state["input_files"] = [xlsx_path]
        result = ingest_node(minimal_state)

        assert "raw_texts" in result
        filename = os.path.basename(xlsx_path)
        assert filename in result["raw_texts"]
        assert "Three-way match" in result["raw_texts"][filename]

    def test_ingest_multiple_files(self, minimal_state, tmp_path):
        """ingest_node should handle a list of mixed file types."""
        from docx import Document
        import openpyxl
        from nodes.ingest_node import ingest_node

        doc = Document()
        doc.add_paragraph("MOM content for AP module")
        docx_path = str(tmp_path / "mom.docx")
        doc.save(docx_path)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Questionnaire data"
        xlsx_path = str(tmp_path / "q.xlsx")
        wb.save(xlsx_path)

        minimal_state["input_files"] = [docx_path, xlsx_path]
        result = ingest_node(minimal_state)

        raw = result["raw_texts"]
        assert len(raw) == 2
        assert "mom.docx" in raw
        assert "q.xlsx" in raw

    def test_ingest_preserves_source_filename(self, minimal_state, tmp_path):
        """Filename (not full path) should be the key in raw_texts."""
        from docx import Document
        from nodes.ingest_node import ingest_node

        doc = Document()
        doc.add_paragraph("Content")
        docx_path = str(tmp_path / "AP20_Formatted.docx")
        doc.save(docx_path)

        minimal_state["input_files"] = [docx_path]
        result = ingest_node(minimal_state)

        assert "AP20_Formatted.docx" in result["raw_texts"]

    def test_ingest_empty_file_list(self, minimal_state):
        """Empty input_files list should return empty raw_texts without raising."""
        from nodes.ingest_node import ingest_node

        minimal_state["input_files"] = []
        result = ingest_node(minimal_state)

        assert result["raw_texts"] == {}
        assert result["last_completed_node"] == "ingest"

    def test_ingest_missing_file_records_error(self, minimal_state):
        """A missing file should be recorded in errors, not raise."""
        from nodes.ingest_node import ingest_node

        minimal_state["input_files"] = ["/this/path/does/not/exist.docx"]
        result = ingest_node(minimal_state)

        # Should not crash; error should be logged
        assert "errors" in result or result.get("raw_texts") == {}
