"""
RD.011 Agent — Word document builder.

Constructs the full RD.011 Future Process Model document using python-docx.
No template rendering (no docxtpl). The style reference .docx is used only
to import paragraph styles at startup.
"""

from __future__ import annotations

import datetime
import logging
import os
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import DIAGRAM_WIDTH_INCHES, TEMPLATE_PATH, WORD_STYLES
from models.schemas import (
    DocumentPlan,
    IntroContent,
    IssueReport,
    JournalEntry,
    SectionContent,
    SectionPlan,
)

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Builds an RD.011 Word document from structured data."""

    def __init__(self, style_ref_path: str = TEMPLATE_PATH):
        """
        Initialise a new blank python-docx Document.

        Attempts to use *style_ref_path* as the base document to inherit
        its paragraph/table styles.  If the file does not exist, creates
        a fully blank document and relies on python-docx built-in styles.
        """
        if os.path.exists(style_ref_path):
            try:
                self.doc = Document(style_ref_path)
                # Clear all existing body content from the style reference,
                # but preserve the final w:sectPr (section properties) that
                # defines page layout (margins, width, orientation).  Without
                # it, add_table() fails with IndexError.
                body = self.doc.element.body
                sectPr_tag = qn("w:sectPr")
                for element in list(body):
                    if element.tag != sectPr_tag:
                        body.remove(element)
                logger.info("Loaded styles from: %s", style_ref_path)
            except Exception as exc:
                logger.warning(
                    "Failed to load style reference %s: %s — using blank document",
                    style_ref_path,
                    exc,
                )
                self.doc = Document()
        else:
            logger.info(
                "Style reference not found at %s — using blank document",
                style_ref_path,
            )
            self.doc = Document()

        self._available_styles = {s.name for s in self.doc.styles}

    # ── Helpers ───────────────────────────────────────────────────────────

    def _style_name(self, key: str) -> Optional[str]:
        """Resolve a style key to a style name if available."""
        name = WORD_STYLES.get(key, key)
        if name in self._available_styles:
            return name
        return None

    def _add_paragraph(self, text: str, style_key: str = "body_text", **kwargs):
        """Add a paragraph with the given style, falling back to Normal."""
        style = self._style_name(style_key)
        p = self.doc.add_paragraph(text, style=style)
        if kwargs.get("bold"):
            for run in p.runs:
                run.bold = True
        if kwargs.get("italic"):
            for run in p.runs:
                run.italic = True
        if kwargs.get("alignment"):
            p.alignment = kwargs["alignment"]
        return p

    def _add_heading(self, text: str, level: int = 1):
        """Add a heading at the specified level."""
        return self.doc.add_heading(text, level=level)

    def _add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    def _apply_table_style(self, table) -> None:
        """Apply a native Word table style.

        'Table Grid' is tried first because it guarantees visible black borders
        in every Word/LibreOffice version.  Falls back to accent styles, then
        applies table-level XML borders as a fallback.
        """
        for name in ("Table Grid", "Light Shading Accent 1", "Light List Accent 1"):
            try:
                table.style = name
                return
            except KeyError:
                continue

        # No built-in table style available — apply table-level XML borders
        logger.debug("No table style available; applying table-level XML borders")
        self._apply_table_borders(table)

    def _apply_table_borders(self, table) -> None:
        """Apply table-level XML borders (tcBorders) for guaranteed visibility.

        This applies borders at the table properties level so they appear
        consistently in all Word versions.
        """
        try:
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)

            # Create table borders element (tblBorders)
            tblBorders = OxmlElement("w:tblBorders")

            # Border attributes: val="single" sz="12" space="0" color="000000"
            border_attrs = {
                qn("w:val"): "single",
                qn("w:sz"): "12",
                qn("w:space"): "0",
                qn("w:color"): "000000",
            }

            # Apply to all border types: top, left, bottom, right, insideH, insideV
            for border_type in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = OxmlElement(f"w:{border_type}")
                for attr_name, attr_value in border_attrs.items():
                    border_el.set(attr_name, attr_value)
                tblBorders.append(border_el)

            # Remove any existing tblBorders and add the new one
            existing_tblBorders = tblPr.find(qn("w:tblBorders"))
            if existing_tblBorders is not None:
                tblPr.remove(existing_tblBorders)
            tblPr.append(tblBorders)

            logger.debug("Applied table-level XML borders")
        except Exception as exc:
            logger.warning("Failed to apply table-level borders: %s", exc)


    def _add_table_header_row(self, table, columns: List[str]) -> None:
        """Add a header row with TableHeading style and bold text."""
        header_cells = table.rows[0].cells
        style_name = self._style_name("table_heading")

        for idx, col_text in enumerate(columns):
            cell = header_cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.text = col_text
            if style_name:
                p.style = self.doc.styles[style_name]
            run = p.runs[0] if p.runs else p.add_run(col_text)
            run.bold = True
            run.font.size = Pt(9)

    def _add_table_data_row(self, table, values: List[str]) -> None:
        """Add a data row with TableText style."""
        row = table.add_row()
        style_name = self._style_name("table_text")

        for idx, val in enumerate(values):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.text = val
            if style_name:
                p.style = self.doc.styles[style_name]
            for run in p.runs:
                run.font.size = Pt(9)

    # ── Cover Page ────────────────────────────────────────────────────────

    def build_cover_page(self, plan: DocumentPlan) -> None:
        """
        Add cover page content.

        - "RD.011 FUTURE PROCESS MODEL" as large centered title
        - "ORACLE FINANCE CLOUD" subtitle
        - Client name (large, bold)
        - Author, Creation Date, Document Ref, Version as a 2-col table
        """
        # Main title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(120)
        run = p.add_run("RD.011 FUTURE PROCESS MODEL")
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ORACLE FINANCE CLOUD")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

        # Client name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(40)
        run = p.add_run(plan.client_name)
        run.font.size = Pt(22)
        run.bold = True

        # Project name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(plan.project_name)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Metadata table
        p = self.doc.add_paragraph()
        p.space_before = Pt(60)

        table = self.doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Compute creation_date as today's date (ISO format)
        creation_date = datetime.date.today().strftime("%Y-%m-%d")

        meta_data = [
            ("Author", plan.author),
            ("Creation Date", creation_date),
            ("Document Reference", plan.document_ref),
            ("Version", plan.version),
        ]
        for i, (label, value) in enumerate(meta_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)
                runs = cell.paragraphs[0].runs
                if runs:
                    runs[0].bold = (cell == row.cells[0])

        self._apply_table_style(table)
        self._add_page_break()

    def build_table_of_contents(self) -> None:
        """Insert a Word Table of Contents field."""
        self._add_heading("Table of Contents", level=1)
        p = self.doc.add_paragraph()
        run = p.add_run()

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_sep)
        run._r.append(fld_char_end)

        self._add_page_break()

    # ── Document Control ──────────────────────────────────────────────────

    def build_document_control(self, plan: DocumentPlan) -> None:
        """
        Add Document Control section with change record, reviewers,
        and approvers tables.
        """
        self._add_heading("Document Control", level=1)

        # Change Record
        self._add_heading("Change Record", level=2)
        table = self.doc.add_table(rows=1, cols=4)
        self._add_table_header_row(table, ["Date", "Author", "Version", "Change Reference"])
        # Use today's date (ISO format) for the change record
        creation_date = datetime.date.today().strftime("%Y-%m-%d")
        self._add_table_data_row(table, [
            creation_date, plan.author, plan.version, "Initial draft"
        ])
        self._apply_table_style(table)

        self.doc.add_paragraph()

        # Reviewers
        self._add_heading("Reviewers", level=2)
        table = self.doc.add_table(rows=1, cols=3)
        self._add_table_header_row(table, ["No.", "Name", "Position"])
        for i in range(3):
            self._add_table_data_row(table, [str(i + 1), "", ""])
        self._apply_table_style(table)

        self.doc.add_paragraph()

        # Approvers
        self._add_heading("Approvers", level=2)
        table = self.doc.add_table(rows=1, cols=4)
        self._add_table_header_row(
            table, ["Name", "Position", "Date of Approval", "Signature"]
        )
        for _ in range(5):
            self._add_table_data_row(table, ["", "", "", ""])
        self._apply_table_style(table)

        self._add_page_break()

    # ── Introduction ──────────────────────────────────────────────────────

    def build_introduction(self, intro: IntroContent, plan: DocumentPlan) -> None:
        """
        Add Introduction section with intro paragraphs, document
        organization, and process symbols explanation.
        """
        self._add_heading("Introduction", level=1)

        for para_text in intro.introduction_paragraphs:
            self._add_paragraph(para_text)

        self._add_heading("How This Document Is Organized", level=2)
        self._add_paragraph(intro.how_organized_text)

        # Process Symbols removed per client output requirements

    # ── Enterprise Structure ──────────────────────────────────────────────

    def build_enterprise_section(self, intro: IntroContent, plan: DocumentPlan) -> None:
        """
        Add the enterprise structure section: context paragraphs, ledger
        facts, and Chart of Accounts description.
        """
        self._add_heading(f"{plan.client_name} Enterprise Structure", level=1)

        for para_text in intro.enterprise_context_paragraphs:
            self._add_paragraph(para_text)

        self._add_heading("Ledger", level=2)
        for fact in intro.ledger_facts:
            style_name = self._style_name("bullet")
            if style_name:
                self.doc.add_paragraph(fact, style=style_name)
            else:
                self._add_paragraph(f"- {fact}")

        self._add_heading("Chart of Accounts Structure", level=2)
        self._add_paragraph(intro.coa_description)

        self._add_page_break()

    # ── Module Chapter ────────────────────────────────────────────────────

    def build_module_chapter(
        self,
        section_plan: SectionPlan,
        sections: Dict[str, SectionContent],
        diagram_registry: Dict[str, str],
        chapter_number: int = 0,
    ) -> None:
        """
        Add a full module chapter with process outline table and
        per-process detail sections.
        """
        # Module heading
        if chapter_number > 0:
            heading_text = f"{chapter_number}. {section_plan.module_name} \u2014 List of Processes"
        else:
            heading_text = f"{section_plan.module_name} \u2014 List of Processes"
        self._add_heading(heading_text, level=1)

        # Module intro
        self._add_paragraph(section_plan.module_intro)

        # Process Outline Table
        table = self.doc.add_table(rows=1, cols=4)
        self._add_table_header_row(
            table, ["ID", "Process Name", "Process Description", "Output"]
        )
        for proc in section_plan.processes:
            self._add_table_data_row(table, [
                proc.process_id,
                proc.process_name,
                proc.process_description,
                proc.output,
            ])
        self._apply_table_style(table)

        self.doc.add_paragraph()

        # Per-process detail sections
        written_ids: set[str] = set()
        for idx, proc in enumerate(section_plan.processes):
            pid = proc.process_id
            if pid in written_ids:
                logger.warning("Duplicate process_id %s - skipping", pid)
                continue
            written_ids.add(pid)

            section_key = f"{section_plan.section_id}.{proc.process_id}"
            # Also try just process_id as key
            content = sections.get(section_key) or sections.get(proc.process_id)

            if content is None:
                self._add_heading(f"{proc.process_id} {proc.process_name}", level=2)
                self._add_paragraph(
                    f"[Content not generated for {proc.process_id}]", italic=True
                )
                continue

            # Ensure content is a SectionContent object
            if isinstance(content, dict):
                content = SectionContent.model_validate(content)

            # Process heading
            self._add_heading(f"{proc.process_id} {content.process_name}", level=2)

            # Narrative
            self._add_heading("Narrative", level=3)
            # Split narrative into paragraphs on double newlines
            for para in content.narrative.split("\n\n"):
                para = para.strip()
                if para:
                    self._add_paragraph(para)

            # Process Diagram
            self._add_heading("Process Diagram", level=3)
            png_path = diagram_registry.get(proc.process_id)
            if png_path and os.path.exists(png_path):
                try:
                    self.doc.add_picture(png_path, width=Inches(DIAGRAM_WIDTH_INCHES))
                except Exception as exc:
                    logger.warning("Could not embed diagram %s: %s", png_path, exc)
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"[Diagram not embedded: {exc}]")
                    run.italic = True
            else:
                p = self.doc.add_paragraph()
                run = p.add_run(
                    "[Diagram not available — install Graphviz (dot) for rendered diagrams]"
                )
                run.italic = True

            # Process Step Catalog
            self._add_heading("Process Step Catalog", level=3)
            if not content.process_steps:
                self._add_paragraph("N/A")
            else:
                table = self.doc.add_table(rows=1, cols=4)
                self._add_table_header_row(
                    table,
                    ["ID", "Action", "Description", "Business Actor"],
                )
                for step in content.process_steps:
                    self._add_table_data_row(table, [
                        step.step_id,
                        step.action,
                        step.description,
                        step.business_actor,
                    ])
                self._apply_table_style(table)

            # Journal Entries
            self._add_heading("Journal Entries", level=3)
            if not content.journal_entries:
                self._add_paragraph("N/A")
            else:
                for i, je in enumerate(content.journal_entries):
                    # Event label (if present)
                    if je.label:
                        label_p = self._add_paragraph(je.label)
                        for run in label_p.runs:
                            run.bold = True

                    # Debit entry
                    dr_p = self.doc.add_paragraph(f"Dr. {je.debit_account}", style="Normal")
                    dr_p.paragraph_format.left_indent = Inches(0.3)

                    # Credit entry
                    cr_p = self.doc.add_paragraph(f"Cr. {je.credit_account}", style="Normal")
                    cr_p.paragraph_format.left_indent = Inches(0.6)

                    # Optional amount label (displayed as secondary text)
                    if je.amount_label:
                        amt_p = self.doc.add_paragraph(f"({je.amount_label})", style="Normal")
                        amt_p.paragraph_format.left_indent = Inches(0.9)

                    # Blank line between entries (except after last)
                    if i < len(content.journal_entries) - 1:
                        self.doc.add_paragraph()

            self.doc.add_paragraph()

            # Key Requirements / Highlights
            self._add_heading("Key Requirements / Highlights", level=3)
            nl_style = self._style_name("number_list")
            for i, req in enumerate(content.key_requirements, start=1):
                if nl_style:
                    self.doc.add_paragraph(req, style=nl_style)
                else:
                    self._add_paragraph(f"{i}. {req}")

            # Page break between processes for clear separation
            if idx < len(section_plan.processes) - 1:
                self._add_page_break()

        # Page break after last process in module
        self._add_page_break()

    # ── Open and Closed Issues ────────────────────────────────────────────

    def build_issues_section(
        self,
        open_issues: Optional[List[dict]] = None,
        closed_issues: Optional[List[dict]] = None,
    ) -> None:
        """
        Add the Open and Closed Issues section with tabular layouts.
        """
        open_issues = open_issues or []
        closed_issues = closed_issues or []

        self._add_heading("Open and Closed Issues", level=1)

        # Open Issues
        self._add_heading("Open Issues", level=2)
        issue_cols = [
            "ID", "Issue", "Resolution", "Responsibility",
            "Target Date", "Impact Date",
        ]
        table = self.doc.add_table(rows=1, cols=6)
        self._add_table_header_row(table, issue_cols)

        if open_issues:
            for issue in open_issues:
                self._add_table_data_row(table, [
                    str(issue.get("id", "")),
                    issue.get("issue", ""),
                    issue.get("resolution", ""),
                    issue.get("responsibility", ""),
                    issue.get("target_date", ""),
                    issue.get("impact_date", ""),
                ])
        else:
            for i in range(3):
                self._add_table_data_row(table, [str(i + 1), "", "", "", "", ""])
        self._apply_table_style(table)

        self.doc.add_paragraph()

        # Closed Issues
        self._add_heading("Closed Issues", level=2)
        table = self.doc.add_table(rows=1, cols=6)
        self._add_table_header_row(table, issue_cols)

        if closed_issues:
            for issue in closed_issues:
                self._add_table_data_row(table, [
                    str(issue.get("id", "")),
                    issue.get("issue", ""),
                    issue.get("resolution", ""),
                    issue.get("responsibility", ""),
                    issue.get("target_date", ""),
                    issue.get("impact_date", ""),
                ])
        else:
            for i in range(3):
                self._add_table_data_row(table, [str(i + 1), "", "", "", "", ""])
        self._apply_table_style(table)

    # ── Diagram Appendix ──────────────────────────────────────────────────

    def build_diagram_appendix(self, generated_sections: Dict[str, dict]) -> None:
        """
        Add an appendix containing the raw Graphviz DOT source for every
        diagram to allow for manual consultant adjustments.
        """
        self._add_page_break()
        self._add_heading("Appendix: Graphviz DOT Diagram Source Code", level=1)
        self._add_paragraph(
            "This appendix contains the raw Graphviz DOT source code for all "
            "process flow diagrams. If a diagram requires adjustments (e.g., adding a step or changing a role), "
            "you do not need to know how to code! Just follow these simple steps:"
        )
        self._add_paragraph("1. Copy the block of code below for the diagram you want to change.")
        self._add_paragraph("2. Go to an online Graphviz editor like: https://dreampuf.github.io/GraphvizOnline/")
        self._add_paragraph("3. Paste the code into the left side of the screen. You will see the diagram appear on the right.")
        self._add_paragraph("4. Modify the plain text inside the code to change step names or business actors. The diagram will update live.")
        self._add_paragraph("5. Download the updated diagram as a PNG image and paste it directly into this Word document to replace the old one!")

        for section_key, section_data in generated_sections.items():
            if isinstance(section_data, dict):
                process_id = section_data.get("process_id", section_key)
                process_name = section_data.get("process_name", "")
                diagram_code = section_data.get("diagram_code", "")
            else:
                process_id = section_data.process_id
                process_name = section_data.process_name
                diagram_code = section_data.diagram_code

            if not diagram_code:
                continue

            self._add_heading(f"{process_id} {process_name}", level=2)

            # Add DOT source in a monospace-style paragraph
            p = self.doc.add_paragraph()
            run = p.add_run(diagram_code)
            run.font.size = Pt(8)
            run.font.name = "Consolas"
            # Set font in underlying XML for Consolas
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rFonts.set(qn("w:cs"), "Consolas")
            rPr.append(rFonts)

    # ── Save ──────────────────────────────────────────────────────────────

    def save(self, output_path: str) -> str:
        """
        Save the document to *output_path*.

        Creates the output directory if it does not exist.

        Returns
        -------
        str
            The *output_path* written to.
        """
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        self.doc.save(output_path)
        logger.info("Document saved: %s", output_path)
        return output_path
