"""
Extract a lightweight outline (headings + key markers) from a .docx file.

This helps answer questions like:
- Does this example follow the same section naming as our template?
- Where does the doc place Enterprise Structure / module chapters / process subsections?

Run:
  conda run -n rd011-env python tools/docx_outline.py templates/RD011_TEMPLATE.docx
  conda run -n rd011-env python tools/docx_outline.py rag/examples/<file>.docx
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


@dataclass(frozen=True)
class Heading:
    level: int
    text: str


def _heading_level(paragraph) -> int | None:
    """
    Best-effort heading detection.

    - Prefer style name "Heading N"
    - Fallback to outline level in the XML when present
    """
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            pass

    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        outlineLvl = pPr.find(qn("w:outlineLvl"))
        if outlineLvl is not None:
            val = outlineLvl.get(qn("w:val"))
            if val is not None:
                try:
                    return int(val) + 1
                except ValueError:
                    return None
    return None


def _norm(text: str) -> str:
    return " ".join((text or "").strip().lower().split())


_MARKERS = [
    "Document Control",
    "Change Record",
    "Reviewers",
    "Introduction",
    "Enterprise Structure",
    "Business Architecture",
    "Event Catalog",
    "Business Actors",
    "Process Outline",
    "Process Diagram",
    "Process Step Catalog",
    "Journal Entries",
    "Key Requirements",
    "Open Issues",
    "Closed Issues",
]


def extract_outline(path: Path) -> tuple[list[Heading], dict[str, int]]:
    doc = Document(str(path))
    headings: list[Heading] = []
    marker_counts = {m: 0 for m in _MARKERS}

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        lvl = _heading_level(para)
        if lvl is not None:
            headings.append(Heading(level=lvl, text=text))

        n = _norm(text)
        for m in _MARKERS:
            if n == _norm(m):
                marker_counts[m] += 1

    # Many RD011 docs put catalogs in tables; scan cell text too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if not t:
                    continue
                n = _norm(t)
                for m in _MARKERS:
                    if n == _norm(m):
                        marker_counts[m] += 1

    return headings, marker_counts


def main(argv: list[str]) -> int:
    # Avoid Windows console encoding issues (cp1252) when documents contain
    # Arabic or special punctuation.
    try:  # Python 3.7+
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    if not argv:
        print("usage: python tools/docx_outline.py <path.docx> [--max N]")
        return 2

    max_items = 180
    args = list(argv)
    if "--max" in args:
        i = args.index("--max")
        try:
            max_items = int(args[i + 1])
        except Exception:
            pass
        del args[i : i + 2]

    path = Path(args[0])
    if not path.exists():
        print(f"not_found: {path}")
        return 2

    headings, markers = extract_outline(path)
    print(f"path={path.resolve()}")
    print(f"headings_count={len(headings)}")

    present = {k: v for k, v in markers.items() if v}
    if present:
        rendered = ", ".join(f"{k}={v}" for k, v in present.items())
        print(f"markers: {rendered}")

    print("\noutline:")
    for idx, h in enumerate(headings[:max_items], 1):
        indent = "  " * max(0, min(h.level, 8) - 1)
        # Trim long headings so output stays readable.
        t = re.sub(r"\s+", " ", h.text).strip()
        print(f"{idx:03d}. {indent}H{h.level}: {t[:160]}")

    if len(headings) > max_items:
        print(f"... truncated (use --max to increase; total={len(headings)})")

    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
