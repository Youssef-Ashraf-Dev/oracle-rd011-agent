"""
Inspect RAG example .docx files for style consistency.

Outputs lightweight stats and a short sample of early lines to help decide:
- Are the example docs written in a consistent consultant voice?
- Are there multiple "style families" that should not be mixed in one prompt?

Run (recommended):
  conda run -n rd011-env python tools/inspect_rag_examples.py
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path


def _safe_lines(doc, limit: int = 8) -> list[str]:
    lines: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        lines.append(t)
        if len(lines) >= limit:
            break
    return lines


_SECTION_MARKERS = [
    "Narrative",
    "Process Step Catalog",
    "Journal Entries",
    "Key Requirements",
    "Key Requirements / Highlights",
    "Process Diagram",
    "Open Issues",
    "Closed Issues",
]


def _norm(text: str) -> str:
    return " ".join((text or "").strip().lower().split())


def main() -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise SystemExit(
            "python-docx is required. Run via the project env, e.g.\n"
            "  conda run -n rd011-env python tools/inspect_rag_examples.py\n"
            f"Error: {exc}"
        )

    examples_dir = Path("rag/examples")
    files = sorted(examples_dir.glob("*.docx"))
    print(f"examples_dir={examples_dir.resolve()}")
    print(f"count={len(files)}")

    style_name_counts: Counter[str] = Counter()

    for path in files:
        doc = Document(str(path))
        nonempty = [p for p in doc.paragraphs if (p.text or "").strip()]
        # Table text matters a lot for RD011 docs (many use tables for catalogs).
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        table_texts.append(t)

        total_chars = sum(len((p.text or "").strip()) for p in nonempty)
        avg_chars = (total_chars / len(nonempty)) if nonempty else 0.0

        heading_paras = 0
        for para in nonempty:
            style = para.style.name if para.style else ""
            style_name_counts[style] += 1
            if style.lower().startswith("heading"):
                heading_paras += 1

        print("\n---", path.name)
        print(
            f"paragraphs={len(nonempty)} headings={heading_paras} "
            f"total_chars={total_chars} avg_para_chars={avg_chars:.1f}"
        )
        for i, line in enumerate(_safe_lines(doc), 1):
            print(f"{i:02d}. {line[:160]}")

        texts = [_norm(p.text) for p in nonempty]
        cell_texts = [_norm(t) for t in table_texts]

        marker_counts = {}
        for marker in _SECTION_MARKERS:
            key = _norm(marker)
            marker_counts[marker] = sum(1 for t in texts if t == key) + sum(1 for t in cell_texts if t == key)

        present = {k: v for k, v in marker_counts.items() if v}
        if present:
            print("section_markers:", ", ".join(f"{k}={v}" for k, v in present.items()))

        # Quick structural proxy: count common "process id" patterns in paragraphs + tables.
        import re

        blob = "\n".join([p.text for p in nonempty] + table_texts)
        dot_ids = re.findall(r"\b[A-Z]{2,3}\.\d{2}\b", blob)
        dash_ids = re.findall(r"\b[A-Z]{2,3}-\d{2}\b", blob)
        compact_ids = re.findall(r"\b[A-Z]{2,3}\d{2}\b", blob)
        print(
            f"process_id_like_count(dot={len(dot_ids)} dash={len(dash_ids)} compact={len(compact_ids)})"
        )

    if style_name_counts:
        most_common = ", ".join(f"{k}={v}" for k, v in style_name_counts.most_common(12))
        print("\nstyle_distribution_top12:", most_common)


if __name__ == "__main__":
    main()
